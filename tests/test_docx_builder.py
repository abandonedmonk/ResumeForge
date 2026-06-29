"""Unit tests for the clean ATS DOCX builder."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document

from app.parsers import docx_builder

_SAMPLE_TEX = r"""
\begin{document}
\begin{center}
    {\Huge \scshape \textbf{Jane R\&D Doe}} \\ \vspace{4pt}
    \href{mailto:jane@example.com}{\ul{jane@example.com}} $|$
    \href{https://github.com/jane}{\ul{GitHub}} \\ \vspace{6pt}
    {Backend engineer headline.}
\end{center}
\section{Experience}
\resumeSubHeadingListStart
  \resumeSubheading
    {Acme \& Co}{Remote}
    {Software Engineer}{2024 -- 2025}
    \resumeItemListStart
      \resumeItem{Built a \textbf{RAG pipeline} over 500 docs}
      \resumeItem{Cut latency by 40\%}
    \resumeItemListEnd
\resumeSubHeadingListEnd
\section{Education}
\resumeSubHeadingListStart
  \resumeSubheading
    {State University}{}
    {B.S. Computer Science}{\textit{2026}}
\resumeSubHeadingListEnd
\end{document}
"""


def test_balanced_handles_nested_braces():
    inner, end = docx_builder._balanced(r"{Built a \textbf{RAG} pipeline}", 0)
    assert inner == r"Built a \textbf{RAG} pipeline"
    assert end == len(r"{Built a \textbf{RAG} pipeline}")


def test_inline_runs_preserve_spacing_and_bold():
    runs = docx_builder._inline_runs(r"Built a \textbf{RAG pipeline} over 500 docs")
    text = "".join(t for t, _ in runs)
    assert text == "Built a RAG pipeline over 500 docs"
    assert ("RAG pipeline", True) in runs


def test_build_docx_roundtrips(tmp_path):
    state = {"final_tex": _SAMPLE_TEX, "generated_headline": "", "output_folder": str(tmp_path)}
    out = tmp_path / "resume.docx"
    docx_builder.build_docx(state, out)  # type: ignore[arg-type]
    assert out.exists()

    doc = Document(str(out))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane R&D Doe" in full
    assert "jane@example.com" in full
    assert "RAG pipeline" in full
    assert "Cut latency by 40%" in full
    assert "State University" in full
    assert "\\" not in full  # no LaTeX commands leaked
    assert "textbf" not in full
    # the metric bullet should carry bold somewhere
    assert any(run.bold for p in doc.paragraphs for run in p.runs)
