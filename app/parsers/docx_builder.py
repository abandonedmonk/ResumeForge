"""Build a clean, ATS-friendly .docx from the tailored resume content.

Not a pixel match of the PDF — a well-structured Word document (name/contact
header, section headings, subheadings, bullet lists) that ATS parsers read
reliably. Works for any template by tokenizing the resume macros in ``final_tex``
and preserving ``\\textbf`` / ``**bold**`` as bold runs.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.agent.state import ResumeState
from app.parsers.latex_parser import parse_latex_resume
from app.utils.config import get_config


def _balanced(text: str, open_index: int) -> tuple[str, int]:
    """Given ``text[open_index] == '{'``, return (inner, index_after_matching_brace)."""
    depth = 0
    for i in range(open_index, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return text[open_index + 1 : i], i + 1
    return text[open_index + 1 :], len(text)


_CMD_RE = re.compile(r"\\(resumeSubheading|resumeProjectHeading|resumeItem|item)\b")
_NEXT_TOKEN_RE = re.compile(
    r"\\(resumeSubheading|resumeProjectHeading|resumeItem|item|resumeItemListEnd"
    r"|resumeSubHeadingListEnd|section|end)\b"
)


def _tokenize(section_tex: str) -> list[tuple[str, list[str]]]:
    """Walk a section's LaTeX and yield ('subheading'|'project'|'item', args) in order."""
    flat = re.sub(r"\s+", " ", section_tex)
    tokens: list[tuple[str, list[str]]] = []
    i, n = 0, len(flat)
    while i < n:
        match = _CMD_RE.search(flat, i)
        if not match:
            break
        cmd = match.group(1)
        cursor = match.end()
        if cmd in ("resumeSubheading", "resumeProjectHeading", "resumeItem"):
            count = {"resumeSubheading": 4, "resumeProjectHeading": 2, "resumeItem": 1}[cmd]
            args: list[str] = []
            for _ in range(count):
                while cursor < n and flat[cursor] == " ":
                    cursor += 1
                if cursor < n and flat[cursor] == "{":
                    content, cursor = _balanced(flat, cursor)
                    args.append(content)
                else:
                    args.append("")
            kind = {"resumeSubheading": "subheading", "resumeProjectHeading": "project", "resumeItem": "item"}[cmd]
            tokens.append((kind, args))
            i = cursor
        else:  # bare \item — content runs until the next macro token
            nxt = _NEXT_TOKEN_RE.search(flat, cursor)
            end = nxt.start() if nxt else n
            tokens.append(("item", [flat[cursor:end]]))
            i = end
    return tokens


def _inline_runs(text: str) -> list[tuple[str, bool]]:
    """Convert a LaTeX fragment to (text, is_bold) runs; strip other commands."""
    text = re.sub(r"\\href\{[^{}]*\}\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\(?:underline|textit|mbox|emph)\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"**\1**", text)
    text = re.sub(r"\\[A-Za-z]+\b", " ", text)  # drop remaining commands (keep their braces' content)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"\$\s*\|\s*\$|\\\\|\$", " ", text)
    for esc, plain in ((r"\&", "&"), (r"\%", "%"), (r"\_", "_"), (r"\#", "#")):
        text = text.replace(esc, plain)

    runs: list[tuple[str, bool]] = []
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            inner = re.sub(r"\s+", " ", part[2:-2]).strip()
            if inner:
                runs.append((inner, True))
        else:
            # Keep boundary spaces (collapse runs of whitespace to one) so adjacent
            # bold/non-bold runs don't get glued together.
            collapsed = re.sub(r"\s+", " ", part)
            if collapsed:
                runs.append((collapsed, False))
    return runs


def _plain(text: str) -> str:
    return re.sub(r"\s+", " ", "".join(t for t, _ in _inline_runs(text))).strip()


def _add_runs(paragraph, text: str) -> None:
    for content, bold in _inline_runs(text):
        run = paragraph.add_run(content)
        run.bold = bold


def _extract_header(tex: str) -> tuple[str, str]:
    """Return (name, contact_line) from the \\begin{center} block."""
    block_match = re.search(r"\\begin\{center\}(.*?)\\end\{center\}", tex, re.DOTALL)
    block = block_match.group(1) if block_match else ""
    name = ""
    bold_match = re.search(r"\\textbf\{([^{}]*)\}", block)
    if bold_match:
        name = _plain(bold_match.group(1))

    pieces: list[str] = []
    cursor = 0
    while True:
        h = block.find(r"\href", cursor)
        if h < 0:
            break
        brace = block.find("{", h)
        if brace < 0:
            break
        url, after = _balanced(block, brace)
        label_brace = block.find("{", after)
        if label_brace < 0:
            cursor = after
            label = ""
        else:
            label, after = _balanced(block, label_brace)
        cursor = after
        clean_label = _plain(label)
        if url.startswith("mailto:"):
            pieces.append(url[len("mailto:") :])
        elif clean_label:
            pieces.append(f"{clean_label}: {url}")
        else:
            pieces.append(url)
    # phone (digits, +, spaces). Drop \vspace{6pt}/\hspace{...} first so "6pt" can't
    # leak a stray digit onto the number.
    phone_src = re.sub(r"\\[vh]space\*?\{[^{}]*\}", " ", block)
    phone = re.search(r"([+]?\d[\d\s().\-]{6,}\d)", _plain(phone_src))
    if phone:
        pieces.append(phone.group(1).strip())
    return name, "  |  ".join(dict.fromkeys(pieces))


def build_docx(state: ResumeState, destination: Path) -> Path:
    tex = state.get("final_tex", "") or ""
    config = get_config()
    document = Document()

    name, contact = _extract_header(tex)
    # The resume's own header is the source of truth; candidate_name is only a fallback.
    name = (name or str(config.get("candidate_name", ""))).strip()
    if name:
        heading = document.add_paragraph()
        run = heading.add_run(name)
        run.bold = True
        run.font.size = Pt(20)
    if contact:
        document.add_paragraph(contact)

    headline = state.get("generated_headline", "")
    if headline:
        document.add_paragraph(_plain(headline))

    for section_name, section in parse_latex_resume(tex).items():
        if section_name.lower() == "document":
            continue
        document.add_heading(section_name, level=1)
        for kind, args in _tokenize(str(section.get("raw_tex", ""))):
            if kind == "subheading":
                org, loc, role, dates = (args + ["", "", "", ""])[:4]
                line = document.add_paragraph()
                _add_runs(line, f"\\textbf{{{_plain(org)}}}")
                tail = " — ".join(p for p in (_plain(role), _plain(dates), _plain(loc)) if p)
                if tail:
                    line.add_run(f"  ({tail})")
            elif kind == "project":
                title = document.add_paragraph()
                _add_runs(title, f"\\textbf{{{_plain(args[0])}}}")
                if len(args) > 1 and _plain(args[1]):
                    title.add_run(f"  ({_plain(args[1])})")
            else:  # item
                text = _plain(args[0]) if args else ""
                if text:
                    bullet = document.add_paragraph(style="List Bullet")
                    _add_runs(bullet, args[0])

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return destination
